#!/usr/bin/env python3
"""生成 HappyMac 雷达信号表征学术报告 (Word/PDF 可打印版)"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ANALYSIS = Path(__file__).resolve().parent / "analysis"
OUT = Path(__file__).resolve().parent.parent / "docs" / "HappyMac雷达信号表征实验报告.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.4

def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def para(text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for r in p.runs: r.bold = True
    return p

def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.name = 'SimSun'
    r.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_fig(path, width_cm, cap):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    caption(cap)

def add_table(headers, rows, title):
    caption(title)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = str(htxt)
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.text = str(v)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

# ═══════════════════════════════════════════════════════
#  封面
# ═══════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n\n低成本 24GHz 双雷达\n桌面尺度人体感知能力表征\n")
r.font.size = Pt(24); r.bold = True; r.font.name = 'SimHei'

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HappyMac 项目 · 信号表征实验报告")
r.font.size = Pt(14); r.font.name = 'SimHei'

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n2026-08-13\n数据采集与分析：三次独立 Session\n硬件：ESP32-C3 + HLK-LD2450 + HLK-LD2410C + ESP32-S3-EYE")
r.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  摘要
# ═══════════════════════════════════════════════════════
h1("摘要")
para("本研究对两个低成本 24GHz 毫米波雷达（HLK-LD2450 双通道定位雷达、HLK-LD2410C 存在检测雷达）"
     "在桌面尺度（0.6-1.2 m）下感知坐姿人体的能力边界进行了系统性表征。三次独立实验 Session"
     "（每次约 4 分钟、9 个行为阶段）配合摄像头视觉基准（MediaPipe 人脸关键点、MiDaS 单目深度）"
     "交叉验证了雷达各信号通道的有效性。主要发现：（1）LD2450 的 X 通道携带真实横向位置信息"
     "（与视觉基准相关 r≈+0.18，三次可重复），但被静止漂移（σ=128 mm/30s）淹没，仅时间趋势可用；"
     "（2）LD2450 的 Y 通道在阶段统计上符合物理预期，但与单目深度的线性关系未获稳定复现；"
     "（3）LD2410C 的静止能量数值（Es）随身体姿态呈约 2 倍动态范围变化，是有效的姿态弱特征；"
     "其移动能量通道（Em）在桌面尺度基本失效（触发率 2-6%）；头部朝向感知被实验证伪。"
     "结论指向一条明确的 TinyML 路线：以 X 趋势、Y 距离、Es 数值、速度四类信号的时间窗统计特征"
     "训练轻量分类器，并以此为基础开展跨模态蒸馏研究。")

# ═══════════════════════════════════════════════════════
#  1. 引言
# ═══════════════════════════════════════════════════════
h1("1  引言")
h2("1.1  项目背景")
para("HappyMac 是一款吸附于书桌背景板的像素交互小宠物，目标是以「无摄像头、无麦克风」的方式感知用户存在与状态。"
     "感知链路依赖三级传感器：SR602 红外（唤醒哨兵）、LD2410C（微动/存在检测）、LD2450（位置追踪）。"
     "该方案的核心假设是：两个总价不足 50 元的 24GHz 雷达在近距离桌面场景中能够提供足够的行为分类信号。")
h2("1.2  研究问题")
para("（1）LD2450 的 X/Y 坐标是否真实反映人体位置？（2）LD2410C 的能量通道携带什么信息？"
     "（3）哪些信号可用于后续 TinyML 行为分类？")

# ═══════════════════════════════════════════════════════
#  2. 方法
# ═══════════════════════════════════════════════════════
h1("2  实验方法")
h2("2.1  实验协议")
para("每 Session 含 9 个有效行为阶段（自由活动阶段因被试未能及时终止而截断，全部排除）：")
add_table(
    ["阶段", "时长", "行为说明"],
    [["still_30s", "30 s", "静止坐姿"],
     ["slow_sweep", "20 s", "缓慢连续左右摆动（每侧约 5 s）"],
     ["quick_points", "20 s", "快速定点移动：左—停—右—停—中"],
     ["ellipse", "20 s", "上半身椭圆运动（顺/逆时针）"],
     ["fwd_back", "20 s", "前倾—回正—后靠"],
     ["head_only", "20 s", "身体静止，仅头部转动"],
     ["natural_typing", "45 s", "自然打字/浏览"],
     ["natural_reach", "20 s", "伸手取物（手机/水杯）"],
     ["stand_sit", "20 s", "站起—坐下交替"]],
    "表 1  行为阶段协议")
h2("2.2  关键实验控制")
para("• 雷达预热：每次 Session 前雷达静置 ≥3 分钟，排除冷启动漂移期（该现象已在前序实验中识别并固定为实验规范）。\n"
     "• 时间戳对齐：雷达串口数据与摄像头 MJPEG 流均以采集计算机墙钟（time.time()）打戳，配对窗口 150 ms；"
     "延迟扫描确认无残留系统偏移（0 ms 时相关性最优）。\n"
     "• 重复性：完整协议独立重复 3 次（Session 160414 / 165801 / 170609）。\n"
     "• 安装方式：LD2450 天线阵列水平、PCB 竖直，正对受试者，距离约 0.7 m。")
h2("2.3  视觉基准")
para("横向位置基准：MediaPipe Face Landmarker（468 关键点）脸部中心横向归一化坐标 head_x。\n"
     "距离基准：MiDaS_small 单目深度估计（intel-isl/MiDaS），取脸部中心 11×11 邻域中位数。"
     "注意 MiDaS 输出「越近值越大」而雷达 Y「越近值越小」，两者预期负相关。")

# ═══════════════════════════════════════════════════════
#  3. 结果
# ═══════════════════════════════════════════════════════
h1("3  结果")

h2("3.1  LD2450 X 通道：横向位置追踪（弱但可重复）")
add_table(
    ["Session", "r(X)", "对齐样本对数"],
    [["160414", "+0.207", "2164"],
     ["165801", "+0.166", "2052"],
     ["170609", "+0.181", "2015"],
     ["均值", "+0.185", "—"]],
    "表 2  视觉基准 head_x 与雷达 X 的 Pearson 相关（三次 Session）")
para("三次实验符号一致为正，雷达 X 确携带横向位置信息，但相关强度弱——绝对定位不可靠。")
add_table(
    ["行为阶段", "r(X)", "解读"],
    [["slow_sweep", "+0.447", "慢速横向移动时追踪最强"],
     ["head_only", "+0.228", "转头产生 X 响应"],
     ["stand_sit", "+0.229", "站起/坐下伴随 X 变化"],
     ["still_30s", "−0.218", "静止期 X 漂移与头部位置反相关"]],
    "表 3  分阶段 X 相关性（Session 160414）")
para("信号主要集中在预期阶段；静止期反相关表明漂移不跟随头部位置——漂移是 X 通道不可靠的根源。")
add_fig(ANALYSIS / "07_vision_vs_radar_20260813_160414.png", 14,
        "图 1  视觉-雷达联合验证散点（Session 160414）：X 验证 / Y 验证 / Es 朝向验证")

h2("3.2  LD2450 Y 通道：距离测量")
add_table(
    ["行为阶段", "S1", "S2", "S3"],
    [["fwd_back", "−0.534", "+0.015", "−0.208"],
     ["natural_typing", "−0.533", "+0.138", "−0.066"],
     ["slow_sweep", "−0.349", "−0.107", "−0.126"],
     ["ellipse", "+0.374", "+0.472", "+0.324"],
     ["still_30s", "−0.184", "−0.257", "+0.179"]],
    "表 4  MiDaS 深度与雷达 Y 的相关（预期负相关）")
para("Session 1 的前后移动阶段呈现强负相关（r=−0.53），符合物理预期，但 Session 2/3 未复现。"
     "可能原因：单目相对深度在近距场景的尺度不稳定；雷达 Y 测「最近强反射体」而 MiDaS 测「人脸区域深度」，两者物理量不完全同构。"
     "Y 通道在阶段统计层面（表 5，Y σ：stand_sit 438 mm vs head_only 51 mm）仍符合物理预期的响应模式。")
add_fig(ANALYSIS / "10_depth_vs_y_20260813_160414.png", 11,
        "图 2  MiDaS 脸部深度与雷达 Y 散点（Session 160414）")

h2("3.3  阶段统计：三次平均")
add_table(
    ["行为阶段", "X 均值(mm)", "X σ(mm)", "Y 均值(mm)", "Y σ(mm)", "Em%", "Es%"],
    [["still_30s", "+9", "128", "644", "65", "1", "100"],
     ["slow_sweep", "+70", "218", "709", "67", "3", "100"],
     ["quick_points", "+71", "212", "750", "76", "4", "100"],
     ["ellipse", "+77", "158", "710", "81", "3", "100"],
     ["fwd_back", "+52", "86", "730", "98", "6", "100"],
     ["head_only", "−2", "79", "742", "51", "6", "100"],
     ["natural_typing", "−22", "84", "722", "62", "2", "100"],
     ["natural_reach", "+9", "100", "724", "58", "5", "100"],
     ["stand_sit", "+131", "482", "1063", "438", "5", "100"]],
    "表 5  各行为阶段雷达统计（三次 Session 平均）")
add_fig(ANALYSIS / "11_multi_session_stats.png", 15,
        "图 3  X/Y 变异性跨阶段对比（三次 Session 平均，含标准差误差棒）")
para("关键观察：持续运动（慢扫 σ=218 mm、快速定点 σ=212 mm）的 X 变异性约为静止漂移基线（128 mm）的 1.7 倍，"
     "运动信号在统计上可区分于漂移；但静止漂移本身构成约 13 cm 的「噪声地板」，"
     "任何依赖 X 绝对值的单帧分类均不可行，只有时间窗内的趋势模式（斜率、方差）才携带可用信息。")

h2("3.4  LD2410C：Es 数值是姿态特征，Em 失效，朝向感知不存在")
add_table(
    ["行为阶段", "Es 均值 (0-100)"],
    [["natural_typing", "72.8"],
     ["head_only", "69.8"],
     ["still_30s", "63.9"],
     ["slow_sweep", "61.6"],
     ["natural_reach", "60.7"],
     ["quick_points", "59.1"],
     ["ellipse", "55.7"],
     ["fwd_back", "44.2"],
     ["stand_sit", "39.4"]],
    "表 6  LD2410C 静止能量 Es 数值按行为排序（三次平均）")
para("Es 数值呈现近 2 倍动态范围（39-73）：上身稳定的行为（打字、只转头）Es 最高，"
     "身体形态剧变的站起/坐下 Es 最低。Es 数值是有效的身体姿态弱特征。")
para("Em（移动能量）在桌面尺度基本失效：三次平均触发率仅 2-6%（前倾后靠也只有 6%）。"
     "原因：LD2410C 出厂参数面向房间尺度（3-5 m）设计，0.6-1 m 近距离人体「太大、太近」，移动检测门控难以区分。")
add_table(
    ["Session", "r(Es, head_yaw)"],
    [["160414", "−0.000"],
     ["165801", "+0.004"],
     ["170609", "−0.116"]],
    "表 7  Es 与头部朝向（MediaPipe head_yaw）相关（三次 Session）")
para("三次相关系数 |r|<0.12：LD2410C 在桌面场景无法感知头部朝向。"
     "「通过 Es 判断用户是否看向设备」的设计假设被实验证伪。")

# ═══════════════════════════════════════════════════════
#  4. 讨论
# ═══════════════════════════════════════════════════════
h1("4  讨论")
h2("4.1  对 TinyML 管线的影响")
para("（1）可用特征收窄：X 趋势（slope）、Y 距离、Es 数值、速度（V）四类信号有效；"
     "X 绝对值、Em、Es 朝向无效。"
     "（2）模型必须学会「何时信任 X」：静止期漂移与运动期信号在单帧上不可分，"
     "但时间序列模式（持续同向 vs 随机游走）可被滑动窗口统计特征捕获。"
     "（3）Es 数值升格：从「存在检测器」升格为「姿态弱特征」，应纳入特征集并做消融实验验证贡献。")
h2("4.2  与文献的一致性")
para("LD2450 静止目标检测不可靠、多径鬼影现象与 ESPHome/Home Assistant 社区大量报告一致。"
     "本文贡献在于首次给出该雷达桌面场景的定量表征：X 静止漂移 128 mm/30s、运动/静止比 1.7x、"
     "Es 姿态动态范围约 2 倍、Em 触发率 <6%。这些数字此前未见公开数据。")

# ═══════════════════════════════════════════════════════
#  5. 局限
# ═══════════════════════════════════════════════════════
h1("5  局限")
para("（1）n=1 被试，仅单一体型与坐姿；（2）MiDaS 单目深度为相对尺度，无法提供绝对距离验证；"
     "（3）三次 Session 时间跨度约 1 小时，未覆盖跨日/跨光照条件；"
     "（4）双雷达同频干扰未做 AB 对照（本实验两雷达同时开启）；"
     "（5）摄像头帧差/关键点基准存在自身误差（MediaPipe 检测率 86-91%）。")

# ═══════════════════════════════════════════════════════
#  6. 结论
# ═══════════════════════════════════════════════════════
h1("6  结论")
para("（1）LD2450 X 轴：横向追踪信号真实存在（r≈+0.18，三次可重复），但被静止漂移淹没，仅趋势可用；"
     "（2）LD2450 Y 轴：阶段统计符合物理预期，与视觉深度基准的线性关系未获稳定复现，需双目/ToF 基准确认；"
     "（3）LD2410C：Es 数值是有效姿态特征，Em 桌面无效，头部朝向感知被证伪；"
     "（4）TinyML 路线成立：以 X 趋势 + Y + Es 数值 + V 为特征集，滑动窗口统计特征 + 轻量分类器；"
     "下一步先以随机森林（RF）验证特征可分性，再蒸馏至 MLP 部署 ESP32-C3。")

# ═══════════════════════════════════════════════════════
#  7. Future Work
# ═══════════════════════════════════════════════════════
h1("7  后续工作")
h2("7.1  短期（TinyML 验证）")
para("• RF 特征可分性验证：用本次三 Session 数据提取滑动窗口特征，以随机森林评估 4-5 类行为"
     "（离场/静止/左右移动/前后移动/站起坐下）的分类准确率，作为 TinyML 的可行性闸门。\n"
     "• 特征消融：分别移除 X 趋势、Es 数值、速度，量化各通道边际贡献。\n"
     "• 阈值规则 baseline：实现手写规则分类器，与 ML 方案定量对比。")
h2("7.2  中期（跨模态蒸馏）")
para("• 摄像头-雷达硬件同步：以 GPIO 线连接 C3 与 S3，实现微秒级时间戳同步，"
     "重新验证 Y 通道与视觉深度的线性关系。\n"
     "• 视觉 Teacher 训练：以 MediaPipe 关键点训练「桌面注意力状态」分类器（天花板实验），"
     "确定 Student 的理论精度上限。\n"
     "• TinyML 部署：训练 2 层 MLP（<2 KB 参数），INT8 量化后部署 ESP32-C3，"
     "实测推理延迟与量化损失。")
h2("7.3  长期（学术方向）")
para("• Paper 1（信号表征）：本文数据扩充后投稿传感系统会议——首个低成本 24GHz 雷达桌面场景定量表征。\n"
     "• Paper 2（跨模态蒸馏）：视觉 Teacher → 雷达 Student 的蒸馏管线，"
     "核心卖点为「信息损失分解：模态 gap vs 量化损失」。\n"
     "• Paper 3（交互设计）：不完美感知的交互设计框架（Attention-Oriented Sensing），"
     "以 HappyMac 概率表情状态机为案例。\n"
     "• 硬件改进：VL53L1X ToF 三角定位补充 X 精度；双雷达分时复用（MOSFET 门控）量化同频干扰。")

# ═══════════════════════════════════════════════════════
#  附录
# ═══════════════════════════════════════════════════════
h1("附录")

import json as _json

h2("附录 A  每 Session 每阶段完整统计")
appendix = _json.load(open(ANALYSIS / "appendix_per_session.json"))
# 只保留三次正式 Session
official = ['20260813_160414', '20260813_165801', '20260813_170609']
act_names = {
    'still_30s': '静止', 'slow_sweep': '慢速左右摆', 'quick_points': '快速定点',
    'ellipse': '椭圆运动', 'fwd_back': '前倾后靠', 'head_only': '只转头',
    'natural_typing': '自然打字', 'natural_reach': '伸手取物', 'stand_sit': '站起坐下',
}
for ts in official:
    acts = appendix.get(ts, [])
    if not acts: continue
    rows = []
    for a in acts:
        rows.append([
            act_names.get(a['action'], a['action']),
            str(a['n']),
            f"{a['x_mean']:+.0f}",
            f"{a['x_std']:.0f}",
            f"{a['y_mean']:.0f}",
            f"{a['y_std']:.0f}",
            f"{a['em_pct']:.0f}%",
            f"{a['es_pct']:.0f}%",
        ])
    add_table(
        ["行为阶段", "帧数", "X均值(mm)", "Xσ(mm)", "Y均值(mm)", "Yσ(mm)", "Em%", "Es%"],
        rows,
        f"表 A.{official.index(ts)+1}  Session {ts} 逐阶段统计")

h2("附录 B  分阶段视觉相关性（Session 160414）")
add_table(
    ["行为阶段", "X r", "Y r"],
    [["ellipse", "+0.150", "+0.273"],
     ["fwd_back", "−0.124", "−0.361"],
     ["head_only", "+0.228", "−0.061"],
     ["natural_reach", "+0.006", "−0.098"],
     ["natural_typing", "−0.163", "+0.116"],
     ["quick_points", "+0.098", "−0.463"],
     ["slow_sweep", "+0.447", "+0.110"],
     ["stand_sit", "+0.229", "+0.356"],
     ["still_30s", "−0.218", "+0.055"]],
    "表 B.1  MediaPipe 基准分阶段相关（X=横向位置，Y=深度 proxy z）")

h2("附录 C  MiDaS 深度验证完整结果")
add_table(
    ["行为阶段", "S1 (160414)", "S2 (165801)", "S3 (170609)"],
    [["ellipse", "+0.374", "+0.472", "+0.324"],
     ["fwd_back", "−0.534", "+0.015", "−0.208"],
     ["head_only", "−0.266", "−0.025", "−0.284"],
     ["natural_reach", "+0.159", "−0.074", "−0.151"],
     ["natural_typing", "−0.533", "+0.138", "−0.066"],
     ["quick_points", "+0.150", "−0.249", "−0.222"],
     ["slow_sweep", "−0.349", "−0.107", "−0.126"],
     ["stand_sit", "−0.093", "−0.110", "+0.265"],
     ["still_30s", "−0.184", "−0.257", "+0.179"],
     ["全部合并", "−0.161", "+0.154", "−0.028"]],
    "表 C.1  MiDaS 脸部深度与雷达 Y 相关（预期负相关）")

h2("附录 D  实验硬件与软件环境")
add_table(
    ["组件", "规格"],
    [["主控", "ESP32-C3 SuperMini（RISC-V 160MHz，400KB SRAM）"],
     ["位置雷达", "HLK-LD2450，1T2R FMCW，24GHz，天线基线 ~6.25mm，输出 10Hz"],
     ["存在雷达", "HLK-LD2410C，24GHz FMCW，9 距离门，0.75m/门"],
     ["摄像头", "ESP32-S3-EYE（AP 模式 MJPEG 推流，10fps 640×480）"],
     ["视觉基准 1", "MediaPipe Face Landmarker（468 关键点，86-91% 检测率）"],
     ["视觉基准 2", "MiDaS_small 单目深度（intel-isl/MiDaS，PyTorch 1.x，CPU 推理）"],
     ["时间戳", "采集计算机墙钟 time.time()，配对窗口 150ms，无残留延迟"]],
    "表 D.1  实验环境")

h2("附录 E  数据分析管线说明")
para("1. collect_session.py：SYNC 握手 + 3 分钟雷达预热 + 逐帧墙钟打戳（雷达 CSV + 视频帧时间戳 vidts.csv 独立保存）。\n"
     "2. analyze_session.py：雷达侧 6 图（时间序列/分布/静止对比/漂移/能量散点/科研面板）。\n"
     "3. analyze_vision.py：MediaPipe 逐帧推理 → 与雷达按墙钟最近邻配对 → 三张验证散点。\n"
     "4. analyze_depth.py：MiDaS 每 0.5s 采样 → 脸部 11×11 邻域中位数 → 与雷达 Y 配对。\n"
     "5. analyze_all.py：多 Session 白名单汇总，输出正文均值±标准差与附录逐轮数据。")

doc.save(str(OUT))
print(f"✅ 报告已生成: {OUT}")
